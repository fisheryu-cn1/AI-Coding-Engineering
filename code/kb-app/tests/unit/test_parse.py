"""Unit tests for the M2 parser modules (TXT / MD / HTML / DOCX).

PDF parser is exercised separately in test_parse_pdf.py because it pulls
in pymupdf4llm (heavier dependency).
"""

from __future__ import annotations

from pathlib import Path

import pytest

from kbapp.parse import extension_for, parse_path
from kbapp.parse.base import ParseError
from kbapp.parse.docx import parse_docx
from kbapp.parse.html import parse_html
from kbapp.parse.md import parse_md
from kbapp.parse.txt import parse_txt


def test_extension_for_lower_cases_and_strips_dot(tmp_path: Path) -> None:
    p = tmp_path / "FOO.PDF"
    assert extension_for(p) == "pdf"


def test_parse_unknown_extension_raises(tmp_path: Path) -> None:
    p = tmp_path / "foo.exe"
    p.write_text("nope")
    with pytest.raises(ParseError):
        parse_path(p)


def test_pdf_kwargs_reads_parse_config() -> None:
    """parse.* tunables are wired into the PDF fast path (09 §4 / P2-1)."""
    from kbapp.core.config import Config
    from kbapp.parse.registry import _pdf_kwargs

    cfg = Config.defaults()
    cfg.raw["parse"]["page_char_norm"] = 2000
    cfg.raw["parse"]["pdf_fast_min_coverage"] = 0.5
    cfg.raw["parse"]["pdf_fast_min_headers"] = 7

    kwargs = _pdf_kwargs(cfg)
    assert kwargs["page_char_norm"] == 2000
    assert kwargs["min_coverage"] == 0.5
    assert kwargs["min_headers"] == 7


def test_pdf_kwargs_none_cfg_is_empty() -> None:
    from kbapp.parse.registry import _pdf_kwargs

    assert _pdf_kwargs(None) == {}


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


def test_parse_txt_returns_flat_structure(tmp_path: Path) -> None:
    p = tmp_path / "a.txt"
    p.write_text("para one.\n\npara two.\n\npara three.", encoding="utf-8")
    result, meta = parse_txt(p)
    assert meta.parser == "txt"
    assert meta.format == "txt"
    assert result.structure == "flat"
    assert "para one" in result.full_text
    assert any("para two" in s.text for s in result.sections)


def test_parse_txt_empty_returns_placeholder_section(tmp_path: Path) -> None:
    p = tmp_path / "e.txt"
    p.write_text("", encoding="utf-8")
    result, meta = parse_txt(p)
    # Empty file → 1 placeholder section so downstream stages don't choke.
    assert len(result.sections) == 1
    assert result.sections[0].text == ""
    assert meta.coverage is None


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def test_parse_md_extracts_heading_tree(tmp_path: Path) -> None:
    p = tmp_path / "doc.md"
    p.write_text(
        "# Top\n\nintro text.\n\n## Sub A\n\nbody A.\n\n## Sub B\n\nbody B.\n",
        encoding="utf-8",
    )
    result, meta = parse_md(p)
    assert meta.parser == "markdown-it"
    assert meta.header_count >= 3
    assert result.structure == "tree"
    titles = [s.title for s in result.sections]
    assert "Top" in titles
    assert "Sub A" in titles


def test_parse_md_no_headings_falls_back_to_single_section(tmp_path: Path) -> None:
    p = tmp_path / "no_h.md"
    p.write_text("Just plain content here.\n", encoding="utf-8")
    result, _meta = parse_md(p)
    assert len(result.sections) == 1
    assert result.sections[0].level == 1


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


def test_parse_html_strips_scripts(tmp_path: Path) -> None:
    p = tmp_path / "page.html"
    p.write_text(
        "<html><body><h1>Title</h1><p>hello.</p>"
        "<script>evil()</script><style>css</style>"
        "<p>world.</p></body></html>",
        encoding="utf-8",
    )
    result, meta = parse_html(p)
    assert "evil()" not in result.full_text
    assert "css" not in result.full_text
    assert meta.parser == "trafilatura"


def test_parse_html_empty_returns_no_sections(tmp_path: Path) -> None:
    p = tmp_path / "e.html"
    p.write_text("<html><body></body></html>", encoding="utf-8")
    result, _meta = parse_html(p)
    # Empty body yields at most a degenerate section, but no crash.
    assert isinstance(result.full_text, str)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_parse_docx_extracts_paragraphs(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    p = tmp_path / "doc.docx"
    doc = Document()
    doc.add_heading("Heading 1")
    doc.add_paragraph("First body paragraph.")
    doc.add_paragraph("Second body paragraph.")
    doc.save(p)
    result, meta = parse_docx(p)
    assert meta.parser == "python-docx"
    assert meta.format == "docx"
    assert "First body" in result.full_text
    assert "Second body" in result.full_text
    assert result.structure == "flat"
