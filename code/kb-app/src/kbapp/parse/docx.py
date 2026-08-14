"""DOCX parser via python-docx (用户决策：简化、不依赖 Docling 重型版面分析).

python-docx walks the document body in order; we treat each
``Heading 1/2/3`` paragraph as a section boundary and collect subsequent
``paragraphs`` into the section's text. Tables are inlined (each row → one
line; columns tab-separated) so the chunk stage still sees the content.

``structure`` is ``"flat"`` because DOCX heading detection is best-effort
(no structural enforcement); downstream stages do not rely on tree semantics
for non-MD formats.
"""

from __future__ import annotations

from pathlib import Path

from kbapp.parse.base import ExtractMeta, ParseError, ParseResult, Section

PARSER_NAME = "python-docx"


def parse_docx(path: Path) -> tuple[ParseResult, ExtractMeta]:
    """Parse a ``.docx`` file via python-docx."""
    try:
        from docx import Document  # provided by python-docx (parse dep)
    except ImportError as e:  # pragma: no cover - parse extra missing
        raise ParseError(f"python-docx 未安装（pip install 'kbapp[parse]'）: {e}") from e

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ParseError(f"读取 DOCX 失败：{path} ({e})") from e

    sections: list[Section] = []
    buf: list[str] = []
    current_title = "全文"
    current_level = 1
    start_line = 0
    line_cursor = 0

    def flush(end_line: int) -> None:
        text = "\n".join(buf).strip()
        if not text and not sections:
            return
        sections.append(
            Section(
                section_path=f"§{len(sections) + 1} {current_title}",
                level=current_level,
                title=current_title,
                page_range=f"L{start_line + 1}-L{end_line}",
                text=text,
            )
        )

    for para in doc.paragraphs:
        line_cursor += 1
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            try:
                level = int(style.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            flush(line_cursor - 1)
            current_level = level
            current_title = (para.text or "(无标题)").strip() or "(无标题)"
            buf = []
            start_line = line_cursor
        else:
            t = (para.text or "").strip()
            if t:
                buf.append(t)

    # Tables after the last paragraph (best-effort — python-docx returns
    # them separately). Append as a trailing flat section.
    for tbl in doc.tables:
        line_cursor += 1
        for row in tbl.rows:
            line_cursor += 1
            buf.append("\t".join(cell.text.strip() for cell in row.cells))

    flush(line_cursor)

    if not sections:
        sections.append(
            Section(
                section_path="全文",
                level=1,
                title="全文",
                page_range=None,
                text="",
            )
        )

    full_text = "\n\n".join(s.text for s in sections)
    return (
        ParseResult(
            full_text=full_text,
            sections=sections,
            structure="flat",
            parser=PARSER_NAME,
            warnings=[],
        ),
        ExtractMeta(
            parser=PARSER_NAME,
            format="docx",
            page_count=None,
            header_count=sum(1 for s in sections if s.title != "全文"),
            coverage=None,
        ),
    )


__all__ = ["PARSER_NAME", "parse_docx"]
